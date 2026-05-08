import io
from typing import Dict, Any, List

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side, numbers
from openpyxl.utils import get_column_letter

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_RIGHT, TA_LEFT

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Colors ───────────────────────────────────────────────────────────────────
DARK_BLUE = "1E3A5F"
MID_BLUE = "2E6DA4"
LIGHT_BLUE = "D6E4F0"
HEADER_GRAY = "4A4A4A"
ACCENT = "27AE60"
WHITE = "FFFFFF"
LIGHT_GRAY = "F5F5F5"


# ══════════════════════════════════════════════════════════════════════════════
#  EXCEL EXPORT
# ══════════════════════════════════════════════════════════════════════════════

def export_excel(statements: Dict[str, Any], notes: List[Dict], currency: str) -> bytes:
    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    _write_bs_sheet(wb, statements, currency)
    _write_is_sheet(wb, statements, currency)
    _write_cf_sheet(wb, statements, currency)
    _write_equity_sheet(wb, statements, currency)
    if notes:
        _write_notes_sheet(wb, notes)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _excel_header(ws, title: str, entity: str, period: str, currency: str):
    ws.merge_cells("A1:E1")
    ws["A1"] = entity
    ws["A1"].font = Font(bold=True, size=14, color=WHITE)
    ws["A1"].fill = PatternFill("solid", fgColor=DARK_BLUE)
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 28

    ws.merge_cells("A2:E2")
    ws["A2"] = title
    ws["A2"].font = Font(bold=True, size=12, color=WHITE)
    ws["A2"].fill = PatternFill("solid", fgColor=MID_BLUE)
    ws["A2"].alignment = Alignment(horizontal="center")
    ws.row_dimensions[2].height = 22

    ws.merge_cells("A3:E3")
    ws["A3"] = f"Period Ending: {period}   |   Currency: {currency}"
    ws["A3"].font = Font(italic=True, size=10, color=HEADER_GRAY)
    ws["A3"].alignment = Alignment(horizontal="center")
    ws.row_dimensions[3].height = 18


def _cat_row(ws, row: int, label: str):
    for col in range(1, 6):
        cell = ws.cell(row=row, column=col)
        cell.fill = PatternFill("solid", fgColor=LIGHT_BLUE)
    ws.cell(row=row, column=1).value = label
    ws.cell(row=row, column=1).font = Font(bold=True, size=10)
    ws.row_dimensions[row].height = 16


def _item_row(ws, row: int, label: str, amount: float, indent: int = 2):
    ws.cell(row=row, column=1).value = "  " * indent + label
    ws.cell(row=row, column=1).font = Font(size=10)
    ws.cell(row=row, column=3).value = amount
    ws.cell(row=row, column=3).number_format = "#,##0.00"
    ws.cell(row=row, column=3).alignment = Alignment(horizontal="right")
    if row % 2 == 0:
        for col in range(1, 6):
            ws.cell(row=row, column=col).fill = PatternFill("solid", fgColor=LIGHT_GRAY)


def _total_row(ws, row: int, label: str, amount: float, bold: bool = True, color: str = DARK_BLUE):
    ws.merge_cells(f"A{row}:B{row}")
    ws.cell(row=row, column=1).value = label
    ws.cell(row=row, column=1).font = Font(bold=bold, size=10, color=color)
    ws.cell(row=row, column=4).value = amount
    ws.cell(row=row, column=4).number_format = "#,##0.00"
    ws.cell(row=row, column=4).font = Font(bold=bold, size=10, color=color)
    ws.cell(row=row, column=4).alignment = Alignment(horizontal="right")
    thin = Side(style="thin")
    ws.cell(row=row, column=4).border = Border(top=thin, bottom=thin)
    ws.row_dimensions[row].height = 18


def _set_col_widths(ws):
    ws.column_dimensions["A"].width = 38
    ws.column_dimensions["B"].width = 5
    ws.column_dimensions["C"].width = 16
    ws.column_dimensions["D"].width = 16
    ws.column_dimensions["E"].width = 8


def _write_bs_sheet(wb, statements: Dict, currency: str):
    ws = wb.create_sheet("Balance Sheet")
    bs = statements["balance_sheet"]
    _excel_header(ws, "Balance Sheet", statements["entity_name"], statements["period_end"], currency)
    _set_col_widths(ws)
    r = 5

    for section_name in ["Current Assets", "Non-Current Assets"]:
        if section_name not in bs["sections"]:
            continue
        _cat_row(ws, r, section_name); r += 1
        for item in bs["sections"][section_name]["items"]:
            _item_row(ws, r, item["label"], item["amount"]); r += 1
        _total_row(ws, r, f"Total {section_name}", bs["sections"][section_name]["total"]); r += 2

    _total_row(ws, r, "TOTAL ASSETS", bs["total_assets"], color=ACCENT); r += 2

    ws.cell(row=r, column=1).value = "LIABILITIES AND EQUITY"
    ws.cell(row=r, column=1).font = Font(bold=True, size=11, color=DARK_BLUE)
    r += 1

    for section_name in ["Current Liabilities", "Non-Current Liabilities", "Equity"]:
        if section_name not in bs["sections"]:
            continue
        _cat_row(ws, r, section_name); r += 1
        for item in bs["sections"][section_name]["items"]:
            _item_row(ws, r, item["label"], item["amount"]); r += 1
        _total_row(ws, r, f"Total {section_name}", bs["sections"][section_name]["total"]); r += 2

    _total_row(ws, r, "TOTAL LIABILITIES AND EQUITY", bs["total_liabilities_and_equity"], color=ACCENT); r += 1

    if abs(bs["balance_check"]) > 0.01:
        ws.cell(row=r, column=1).value = f"⚠ Balance Check Difference: {bs['balance_check']:,.2f}"
        ws.cell(row=r, column=1).font = Font(color="FF0000", bold=True)


def _write_is_sheet(wb, statements: Dict, currency: str):
    ws = wb.create_sheet("Income Statement")
    is_ = statements["income_statement"]
    _excel_header(ws, "Income Statement", statements["entity_name"], statements["period_end"], currency)
    _set_col_widths(ws)
    r = 5

    order = ["Revenue", "Cost of Sales", "Operating Expenses", "Finance Costs", "Income Tax"]
    for section_name in order:
        if section_name not in is_["sections"]:
            continue
        _cat_row(ws, r, section_name); r += 1
        for item in is_["sections"][section_name]["items"]:
            _item_row(ws, r, item["label"], item["amount"]); r += 1
        _total_row(ws, r, f"Total {section_name}", is_["sections"][section_name]["total"]); r += 2

    r += 1
    summaries = [
        ("Gross Profit", is_["gross_profit"]),
        ("EBIT", is_["ebit"]),
        ("EBT (Profit before Tax)", is_["ebt"]),
        ("Net Income / (Loss)", is_["net_income"]),
    ]
    for label, val in summaries:
        _total_row(ws, r, label, val, color=ACCENT); r += 1


def _write_cf_sheet(wb, statements: Dict, currency: str):
    ws = wb.create_sheet("Cash Flow")
    cf = statements["cash_flow"]
    _excel_header(ws, "Cash Flow Statement", statements["entity_name"], statements["period_end"], currency)
    _set_col_widths(ws)
    r = 5

    for section_name in ["Operating Activities", "Investing Activities", "Financing Activities"]:
        if section_name not in cf["sections"]:
            continue
        _cat_row(ws, r, section_name); r += 1
        for item in cf["sections"][section_name]["items"]:
            _item_row(ws, r, item["label"], item["amount"]); r += 1
        _total_row(ws, r, f"Net Cash – {section_name}", cf["sections"][section_name]["total"]); r += 2

    _total_row(ws, r, "Net Change in Cash", cf["net_change_in_cash"], color=ACCENT)


def _write_equity_sheet(wb, statements: Dict, currency: str):
    ws = wb.create_sheet("Changes in Equity")
    eq = statements["equity_statement"]
    _excel_header(ws, "Statement of Changes in Equity", statements["entity_name"], statements["period_end"], currency)
    _set_col_widths(ws)
    r = 5

    rows = [
        ("Opening Equity", eq["opening_equity"]),
        ("Net Income for the Period", eq["net_income"]),
        ("Dividends Declared", -abs(eq["dividends"])),
        ("Other Movements", eq["other_movements"]),
        ("Closing Equity", eq["closing_equity"]),
    ]
    for label, val in rows:
        _total_row(ws, r, label, val, bold=(label == "Closing Equity"), color=DARK_BLUE if label != "Closing Equity" else ACCENT)
        r += 1


def _write_notes_sheet(wb, notes: List[Dict]):
    ws = wb.create_sheet("Notes")
    ws.column_dimensions["A"].width = 90
    r = 1
    ws.cell(row=r, column=1).value = "NOTES TO THE FINANCIAL STATEMENTS"
    ws.cell(row=r, column=1).font = Font(bold=True, size=14, color=DARK_BLUE)
    r += 2
    for note in notes:
        ws.cell(row=r, column=1).value = f"Note {note['note_number']}: {note['title']}"
        ws.cell(row=r, column=1).font = Font(bold=True, size=11, color=MID_BLUE)
        r += 1
        import re
        clean = re.sub(r"<[^>]+>", " ", note["body"]).strip()
        ws.cell(row=r, column=1).value = clean
        ws.cell(row=r, column=1).font = Font(size=10)
        ws.cell(row=r, column=1).alignment = Alignment(wrap_text=True)
        ws.row_dimensions[r].height = max(40, min(200, len(clean) // 2))
        r += 2


# ══════════════════════════════════════════════════════════════════════════════
#  PDF EXPORT
# ══════════════════════════════════════════════════════════════════════════════

def export_pdf(statements: Dict[str, Any], notes: List[Dict], currency: str) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle("Title", parent=styles["Title"], fontSize=16, textColor=colors.HexColor(f"#{DARK_BLUE}"), spaceAfter=4)
    subtitle_style = ParagraphStyle("Subtitle", parent=styles["Normal"], fontSize=11, textColor=colors.HexColor(f"#{MID_BLUE}"), spaceAfter=2)
    body_style = ParagraphStyle("Body", parent=styles["Normal"], fontSize=9, spaceAfter=2)
    cat_style = ParagraphStyle("Cat", parent=styles["Normal"], fontSize=10, fontName="Helvetica-Bold", textColor=colors.white, spaceAfter=1)

    story = []

    def section_header(title):
        story.append(Spacer(1, 0.3*cm))
        story.append(Paragraph(title, subtitle_style))
        story.append(Spacer(1, 0.1*cm))

    def add_statement(title, rows_data):
        story.append(Paragraph(title, title_style))
        story.append(Paragraph(f"{statements['entity_name']}   |   Period Ending: {statements['period_end']}   |   {currency}", body_style))
        story.append(Spacer(1, 0.3*cm))

        col_widths = [11*cm, 4*cm, 4*cm]
        table = Table(rows_data, colWidths=col_widths)
        ts = TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{DARK_BLUE}")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, 0), 10),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor(f"#{LIGHT_GRAY}")]),
            ("ALIGN", (1, 0), (-1, -1), "RIGHT"),
            ("FONTSIZE", (0, 1), (-1, -1), 9),
            ("LINEBELOW", (0, 0), (-1, 0), 1, colors.HexColor(f"#{MID_BLUE}")),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ])
        table.setStyle(ts)
        story.append(table)
        story.append(PageBreak())

    # Balance Sheet
    bs = statements["balance_sheet"]
    bs_rows = [["Account", "Amount", ""]]
    for section_name in ["Current Assets", "Non-Current Assets"]:
        if section_name in bs["sections"]:
            bs_rows.append([Paragraph(f"<b>{section_name}</b>", body_style), "", ""])
            for item in bs["sections"][section_name]["items"]:
                bs_rows.append([f"  {item['label']}", f"{item['amount']:,.2f}", ""])
            bs_rows.append([Paragraph(f"<b>Total {section_name}</b>", body_style), "", f"{bs['sections'][section_name]['total']:,.2f}"])
    bs_rows.append([Paragraph("<b>TOTAL ASSETS</b>", body_style), "", f"{bs['total_assets']:,.2f}"])
    bs_rows.append(["", "", ""])
    for section_name in ["Current Liabilities", "Non-Current Liabilities", "Equity"]:
        if section_name in bs["sections"]:
            bs_rows.append([Paragraph(f"<b>{section_name}</b>", body_style), "", ""])
            for item in bs["sections"][section_name]["items"]:
                bs_rows.append([f"  {item['label']}", f"{item['amount']:,.2f}", ""])
            bs_rows.append([Paragraph(f"<b>Total {section_name}</b>", body_style), "", f"{bs['sections'][section_name]['total']:,.2f}"])
    bs_rows.append([Paragraph("<b>TOTAL LIABILITIES AND EQUITY</b>", body_style), "", f"{bs['total_liabilities_and_equity']:,.2f}"])
    add_statement("Balance Sheet", bs_rows)

    # Income Statement
    is_ = statements["income_statement"]
    is_rows = [["Account", "Amount", ""]]
    for section_name in ["Revenue", "Cost of Sales", "Operating Expenses", "Finance Costs", "Income Tax"]:
        if section_name in is_["sections"]:
            is_rows.append([Paragraph(f"<b>{section_name}</b>", body_style), "", ""])
            for item in is_["sections"][section_name]["items"]:
                is_rows.append([f"  {item['label']}", f"{item['amount']:,.2f}", ""])
            is_rows.append([Paragraph(f"<b>Total {section_name}</b>", body_style), "", f"{is_['sections'][section_name]['total']:,.2f}"])
    is_rows.append(["", "", ""])
    is_rows.append([Paragraph("<b>Gross Profit</b>", body_style), "", f"{is_['gross_profit']:,.2f}"])
    is_rows.append([Paragraph("<b>EBIT</b>", body_style), "", f"{is_['ebit']:,.2f}"])
    is_rows.append([Paragraph("<b>Net Income / (Loss)</b>", body_style), "", f"{is_['net_income']:,.2f}"])
    add_statement("Income Statement", is_rows)

    # Cash Flow
    cf = statements["cash_flow"]
    cf_rows = [["Activity", "Amount", ""]]
    for section_name in ["Operating Activities", "Investing Activities", "Financing Activities"]:
        if section_name in cf["sections"]:
            cf_rows.append([Paragraph(f"<b>{section_name}</b>", body_style), "", ""])
            for item in cf["sections"][section_name]["items"]:
                cf_rows.append([f"  {item['label']}", f"{item['amount']:,.2f}", ""])
            cf_rows.append([Paragraph(f"<b>Net Cash – {section_name}</b>", body_style), "", f"{cf['sections'][section_name]['total']:,.2f}"])
    cf_rows.append([Paragraph("<b>Net Change in Cash</b>", body_style), "", f"{cf['net_change_in_cash']:,.2f}"])
    add_statement("Cash Flow Statement", cf_rows)

    # Notes
    if notes:
        story.append(Paragraph("Notes to the Financial Statements", title_style))
        story.append(Spacer(1, 0.3*cm))
        import re
        for note in notes:
            story.append(Paragraph(f"<b>Note {note['note_number']}: {note['title']}</b>", subtitle_style))
            clean = re.sub(r"<[^>]+>", " ", note["body"]).strip()
            story.append(Paragraph(clean, body_style))
            story.append(Spacer(1, 0.3*cm))

    doc.build(story)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════════════════════
#  WORD EXPORT
# ══════════════════════════════════════════════════════════════════════════════

def export_word(statements: Dict[str, Any], notes: List[Dict], currency: str) -> bytes:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    def add_heading(text, level=1):
        p = doc.add_heading(text, level=level)
        if level == 1:
            p.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
        else:
            p.runs[0].font.color.rgb = RGBColor(0x2E, 0x6D, 0xA4)
        return p

    def add_info_line():
        p = doc.add_paragraph()
        run = p.add_run(f"{statements['entity_name']}   |   Period Ending: {statements['period_end']}   |   {currency}")
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
        p.paragraph_format.space_after = Pt(6)

    def add_financial_table(rows):
        t = doc.add_table(rows=0, cols=3)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.LEFT

        for i, row_data in enumerate(rows):
            row = t.add_row()
            for j, val in enumerate(row_data):
                cell = row.cells[j]
                cell.text = str(val)
                p = cell.paragraphs[0]
                run = p.runs[0] if p.runs else p.add_run(str(val))
                run.font.size = Pt(9)
                if j > 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                if i == 0:
                    run.bold = True
                    cell._tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:fill"), DARK_BLUE)
                    cell._tc.tcPr.append(shd)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        t.columns[0].width = Cm(10)
        t.columns[1].width = Cm(3.5)
        t.columns[2].width = Cm(3.5)
        doc.add_paragraph()

    # Balance Sheet
    add_heading("Balance Sheet")
    add_info_line()
    bs = statements["balance_sheet"]
    bs_rows = [["Account", "Amount", "Total"]]
    for section_name in ["Current Assets", "Non-Current Assets"]:
        if section_name in bs["sections"]:
            bs_rows.append([section_name, "", ""])
            for item in bs["sections"][section_name]["items"]:
                bs_rows.append([f"  {item['label']}", f"{item['amount']:,.2f}", ""])
            bs_rows.append([f"Total {section_name}", "", f"{bs['sections'][section_name]['total']:,.2f}"])
    bs_rows.append(["TOTAL ASSETS", "", f"{bs['total_assets']:,.2f}"])
    bs_rows.append(["", "", ""])
    for section_name in ["Current Liabilities", "Non-Current Liabilities", "Equity"]:
        if section_name in bs["sections"]:
            bs_rows.append([section_name, "", ""])
            for item in bs["sections"][section_name]["items"]:
                bs_rows.append([f"  {item['label']}", f"{item['amount']:,.2f}", ""])
            bs_rows.append([f"Total {section_name}", "", f"{bs['sections'][section_name]['total']:,.2f}"])
    bs_rows.append(["TOTAL LIABILITIES AND EQUITY", "", f"{bs['total_liabilities_and_equity']:,.2f}"])
    add_financial_table(bs_rows)

    doc.add_page_break()

    # Income Statement
    add_heading("Income Statement")
    add_info_line()
    is_ = statements["income_statement"]
    is_rows = [["Account", "Amount", "Total"]]
    for section_name in ["Revenue", "Cost of Sales", "Operating Expenses", "Finance Costs", "Income Tax"]:
        if section_name in is_["sections"]:
            is_rows.append([section_name, "", ""])
            for item in is_["sections"][section_name]["items"]:
                is_rows.append([f"  {item['label']}", f"{item['amount']:,.2f}", ""])
            is_rows.append([f"Total {section_name}", "", f"{is_['sections'][section_name]['total']:,.2f}"])
    is_rows.extend([
        ["", "", ""],
        ["Gross Profit", "", f"{is_['gross_profit']:,.2f}"],
        ["EBIT", "", f"{is_['ebit']:,.2f}"],
        ["Net Income / (Loss)", "", f"{is_['net_income']:,.2f}"],
    ])
    add_financial_table(is_rows)

    doc.add_page_break()

    # Cash Flow
    add_heading("Cash Flow Statement")
    add_info_line()
    cf = statements["cash_flow"]
    cf_rows = [["Activity", "Amount", "Total"]]
    for section_name in ["Operating Activities", "Investing Activities", "Financing Activities"]:
        if section_name in cf["sections"]:
            cf_rows.append([section_name, "", ""])
            for item in cf["sections"][section_name]["items"]:
                cf_rows.append([f"  {item['label']}", f"{item['amount']:,.2f}", ""])
            cf_rows.append([f"Net Cash – {section_name}", "", f"{cf['sections'][section_name]['total']:,.2f}"])
    cf_rows.append(["Net Change in Cash", "", f"{cf['net_change_in_cash']:,.2f}"])
    add_financial_table(cf_rows)

    # Notes
    if notes:
        doc.add_page_break()
        add_heading("Notes to the Financial Statements")
        import re
        for note in notes:
            add_heading(f"Note {note['note_number']}: {note['title']}", level=2)
            clean = re.sub(r"<[^>]+>", " ", note["body"]).strip()
            p = doc.add_paragraph(clean)
            p.runs[0].font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
